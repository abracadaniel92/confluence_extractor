"""
Confluence Folder Extractor
============================
Extracts all pages from Confluence folders recursively and exports them to:
- Word documents (.docx)
- Plain text files (.txt)

Each folder gets its own output directory with all pages preserved.

Usage:
    python confluence_folder_extractor.py "https://your-domain.atlassian.net/wiki/spaces/YourSpace/folder/1234567890"
    
    Or multiple folders:
    python confluence_folder_extractor.py "url1" "url2" "url3"
"""

import os
import sys
import re
import requests
import json
import base64
import logging
from pathlib import Path
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# --- Configuration ---
# Load credentials from Tokens_txt.txt (check parent directory first, then current)
TOKENS_FILE = Path(__file__).parent.parent / "Tokens_txt.txt"
if not TOKENS_FILE.exists():
    TOKENS_FILE = Path(__file__).parent / "Tokens_txt.txt"

def load_credentials():
    """Load credentials from Tokens_txt.txt file."""
    credentials = {}
    if not TOKENS_FILE.exists():
        raise FileNotFoundError(f"Credentials file not found: {TOKENS_FILE}")
    
    with open(TOKENS_FILE, 'r') as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith('#') and '=' in line:
                key, value = line.split('=', 1)
                credentials[key.strip()] = value.strip()
    
    CONFLUENCE_BASE_URL = credentials.get('CONFLUENCE_BASE_URL', '')
    CONFLUENCE_API_EMAIL = credentials.get('CONFLUENCE_API_EMAIL', '')
    CONFLUENCE_API_TOKEN = credentials.get('CONFLUENCE_API_TOKEN', '')
    
    if not all([CONFLUENCE_BASE_URL, CONFLUENCE_API_EMAIL, CONFLUENCE_API_TOKEN]):
        raise ValueError("Missing required credentials. Check Tokens_txt.txt file.")
    
    return CONFLUENCE_BASE_URL.rstrip('/'), CONFLUENCE_API_EMAIL, CONFLUENCE_API_TOKEN

CONFLUENCE_BASE_URL, CONFLUENCE_API_EMAIL, CONFLUENCE_API_TOKEN = load_credentials()

# Output base directory (same directory as script)
OUTPUT_BASE_DIR = Path(__file__).parent


# --- UTILITY FUNCTIONS ---
def get_confluence_auth_headers():
    """Generates basic authentication headers for Confluence API."""
    auth_string = f"{CONFLUENCE_API_EMAIL}:{CONFLUENCE_API_TOKEN}"
    encoded_auth = base64.b64encode(auth_string.encode()).decode()
    return {
        "Authorization": f"Basic {encoded_auth}",
        "Content-Type": "application/json",
        "Accept": "application/json"
    }


def extract_text_from_adf(node, include_tables=True):
    """Recursively extracts text from an Atlassian Document Format (ADF) node."""
    if not node or not isinstance(node, dict):
        return ""
    
    node_type = node.get("type", "")
    
    # Handle expand nodes (collapsible sections)
    if node_type == "expand":
        # Extract title if available
        title = ""
        if "attrs" in node and "title" in node["attrs"]:
            title = node["attrs"]["title"]
        
        content_text = ""
        if "content" in node and isinstance(node["content"], list):
            for child in node["content"]:
                child_text = extract_text_from_adf(child, include_tables)
                if child_text:
                    content_text += child_text
        
        if title:
            return f"\n[{title}]\n{content_text}\n"
        return content_text
    
    # Handle tables
    if node_type == "table" and include_tables:
        return extract_table_from_adf(node)
    
    # Handle table rows
    if node_type == "tableRow":
        return ""  # Handled by extract_table_from_adf
    
    # Handle table cells
    if node_type == "tableCell" or node_type == "tableHeader":
        cell_text = ""
        if "content" in node and isinstance(node["content"], list):
            for child in node["content"]:
                child_text = extract_text_from_adf(child, include_tables)
                if child_text:
                    cell_text += child_text + " "
        return cell_text.strip()
    
    # Handle direct text
    if "text" in node:
        text = node["text"]
        # Handle links
        marks = node.get("marks", [])
        for mark in marks:
            if mark.get("type") == "link" and "attrs" in mark:
                url = mark["attrs"].get("href", "")
                if url:
                    return f"{text} [{url}]"
        return text
    
    # Process content list recursively
    text_parts = []
    if "content" in node and isinstance(node["content"], list):
        for child in node["content"]:
            child_text = extract_text_from_adf(child, include_tables)
            if child_text:
                text_parts.append(child_text)
    
    # Add line breaks for certain node types
    if node_type in ["paragraph", "heading"]:
        return "\n".join(text_parts)
    
    return " ".join(text_parts)


def extract_table_from_adf(table_node):
    """Extracts table data from ADF table node and formats it."""
    if not table_node or table_node.get("type") != "table":
        return ""
    
    rows = []
    if "content" in table_node and isinstance(table_node["content"], list):
        for row_node in table_node["content"]:
            if row_node.get("type") == "tableRow":
                cells = []
                if "content" in row_node and isinstance(row_node["content"], list):
                    for cell_node in row_node["content"]:
                        if cell_node.get("type") in ["tableCell", "tableHeader"]:
                            cell_text = ""
                            if "content" in cell_node and isinstance(cell_node["content"], list):
                                for child in cell_node["content"]:
                                    child_text = extract_text_from_adf(child, include_tables=False)
                                    if child_text:
                                        cell_text += child_text + " "
                            cells.append(cell_text.strip())
                if cells:
                    rows.append(cells)
    
    if not rows:
        return ""
    
    # Format as a table
    # Calculate column widths (with minimum width)
    max_cols = max(len(row) for row in rows) if rows else 0
    if max_cols == 0:
        return ""
    
    col_widths = [10] * max_cols  # Minimum width of 10
    
    for row in rows:
        for i, cell in enumerate(row):
            if i < len(col_widths):
                col_widths[i] = max(col_widths[i], min(len(str(cell)) + 2, 80))  # Max 80 chars per column
    
    # Build formatted table
    table_lines = []
    for row_idx, row in enumerate(rows):
        # Pad row to max_cols
        padded_row = row + [""] * (max_cols - len(row))
        
        # Format cells
        formatted_cells = []
        for i, cell in enumerate(padded_row):
            if i < len(col_widths):
                # Truncate if too long, but keep minimum readable
                cell_str = str(cell)
                if len(cell_str) > 80:
                    cell_str = cell_str[:77] + "..."
                formatted_cells.append(cell_str.ljust(min(col_widths[i], 80)))
        
        table_lines.append(" | ".join(formatted_cells))
        
        # Add separator after header row (first row)
        if row_idx == 0 and len(rows) > 1:
            separator = " | ".join(["-" * min(width, 80) for width in col_widths])
            table_lines.append(separator)
    
    return "\n" + "\n".join(table_lines) + "\n\n"


def html_to_plain_text(html_string):
    """Converts HTML string to plain text using BeautifulSoup."""
    if not html_string:
        return ""
    
    try:
        soup = BeautifulSoup(html_string, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        return text
    except Exception as e:
        logging.warning(f"Error parsing HTML: {e}")
        return html_string


def parse_body_content(body_field, extract_plain_text=True):
    """Extracts content from Confluence's body field."""
    if not body_field:
        return ""
    
    # Try storage first (ADF format), then view (HTML), then editor
    storage = body_field.get("storage")
    view = body_field.get("view")
    
    # Prefer ADF storage format for better structure
    if storage:
        if isinstance(storage, dict):
            if "type" in storage or "content" in storage:
                if extract_plain_text:
                    return extract_text_from_adf(storage)
                else:
                    return json.dumps(storage)
            if "value" in storage:
                value = storage["value"]
                if extract_plain_text and isinstance(value, str):
                    if value.strip().startswith('<'):
                        return html_to_plain_text_with_tables(value)
                    return value
                return value
        
        if isinstance(storage, str):
            if extract_plain_text and storage.strip().startswith('<'):
                return html_to_plain_text_with_tables(storage)
            return storage
    
    # Fallback to view (HTML format) which might have better table representation
    if view:
        if isinstance(view, dict) and "value" in view:
            html_content = view["value"]
            if extract_plain_text and isinstance(html_content, str):
                return html_to_plain_text_with_tables(html_content)
        elif isinstance(view, str):
            if extract_plain_text:
                return html_to_plain_text_with_tables(view)
            return view
    
    return str(storage or view or "")


def html_to_plain_text_with_tables(html_string):
    """Converts HTML string to plain text, preserving table structure."""
    if not html_string:
        return ""
    
    try:
        soup = BeautifulSoup(html_string, 'html.parser')
        
        # Extract tables first and replace with formatted text
        tables = soup.find_all('table')
        for table in tables:
            table_text = extract_html_table(table)
            # Replace table with a placeholder, then we'll replace it
            table.replace_with(f"\n{table_text}\n")
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text and clean up whitespace
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        return text
    except Exception as e:
        logging.warning(f"Error parsing HTML: {e}")
        return html_string


def extract_html_table(table_element):
    """Extracts and formats a table from HTML."""
    rows = []
    
    # Find all rows (handle thead, tbody, tfoot)
    for row in table_element.find_all(['tr']):
        cells = []
        for cell in row.find_all(['td', 'th']):
            cell_text = cell.get_text(separator=' ', strip=True)
            cells.append(cell_text)
        if cells:
            rows.append(cells)
    
    if not rows:
        return ""
    
    # Calculate column widths
    max_cols = max(len(row) for row in rows) if rows else 0
    if max_cols == 0:
        return ""
    
    col_widths = [10] * max_cols
    
    for row in rows:
        for i, cell in enumerate(row):
            if i < len(col_widths):
                col_widths[i] = max(col_widths[i], min(len(str(cell)) + 2, 80))
    
    # Build formatted table
    table_lines = []
    for row_idx, row in enumerate(rows):
        padded_row = row + [""] * (max_cols - len(row))
        formatted_cells = []
        for i, cell in enumerate(padded_row):
            if i < len(col_widths):
                cell_str = str(cell)
                if len(cell_str) > 80:
                    cell_str = cell_str[:77] + "..."
                formatted_cells.append(cell_str.ljust(min(col_widths[i], 80)))
        
        table_lines.append(" | ".join(formatted_cells))
        
        # Add separator after header row
        if row_idx == 0 and len(rows) > 1:
            separator = " | ".join(["-" * min(width, 80) for width in col_widths])
            table_lines.append(separator)
    
    return "\n" + "\n".join(table_lines) + "\n"


def sanitize_filename(filename):
    """Sanitize filename to remove invalid characters."""
    # Remove invalid characters for Windows/Linux/Mac
    invalid_chars = '<>:"/\\|?*'
    for char in invalid_chars:
        filename = filename.replace(char, '_')
    # Remove leading/trailing spaces and dots
    filename = filename.strip(' .')
    # Limit length
    if len(filename) > 200:
        filename = filename[:200]
    return filename


# --- CONFLUENCE API FUNCTIONS ---
def fetch_page_by_id(page_id):
    """Fetches a specific Confluence page by its ID with full body content."""
    headers = get_confluence_auth_headers()
    url = f"{CONFLUENCE_BASE_URL}/wiki/rest/api/content/{page_id}"
    
    # Use body.storage to get ADF format which includes all content including tables and expandable sections
    params = {
        "expand": "body.storage,body.view,body.atlas_doc_format,space,version,ancestors"
    }
    
    try:
        response = requests.get(url, headers=headers, params=params)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        logging.error(f"Failed to fetch page {page_id}: {e}")
        if hasattr(e, 'response') and e.response is not None:
            logging.error(f"Response status: {e.response.status_code}")
            logging.error(f"Response body: {e.response.text}")
        return None


def get_folder_info(folder_id):
    """Get folder/page information."""
    page = fetch_page_by_id(folder_id)
    if not page:
        return None, None
    return page.get("title", "Unknown"), page.get("space", {}).get("key", "")


def get_children_recursive(parent_id, limit=50):
    """
    Recursively fetches all children (pages and subfolders) of a parent.
    Returns list of all pages found.
    """
    all_pages = []
    start = 0
    
    headers = get_confluence_auth_headers()
    url = f"{CONFLUENCE_BASE_URL}/wiki/rest/api/content/{parent_id}/child"
    
    while True:
        params = {
            "expand": "body.storage,body.view,space,version,ancestors",
            "limit": limit,
            "start": start
        }
        
        try:
            response = requests.get(url, headers=headers, params=params)
            response.raise_for_status()
            data = response.json()
            
            results = data.get("results", [])
            if not results:
                break
            
            for item in results:
                item_type = item.get("type", "")
                
                # If it's a page
                if item_type == "page":
                    item_id = item.get("id")
                    
                    # Check if this page has children (could be a subfolder)
                    children_check_url = f"{CONFLUENCE_BASE_URL}/wiki/rest/api/content/{item_id}/child"
                    children_check_response = requests.get(
                        children_check_url, 
                        headers=headers, 
                        params={"limit": 1}
                    )
                    
                    has_children = False
                    if children_check_response.status_code == 200:
                        children_data = children_check_response.json()
                        if children_data.get("size", 0) > 0:
                            # This page has children, so it's a subfolder - recurse
                            has_children = True
                            logging.info(f"Found subfolder: {item.get('title', 'N/A')} - recursing...")
                            sub_pages = get_children_recursive(item_id, limit)
                            all_pages.extend(sub_pages)
                    
                    # Fetch full page content with body and add it
                    # (Include folder pages too, as they might contain content)
                    full_page = fetch_page_by_id(item_id)
                    if full_page:
                        all_pages.append(full_page)
                        page_type = "subfolder" if has_children else "page"
                        logging.info(f"Found {page_type}: {item.get('title', 'N/A')}")
                    else:
                        # Fallback: use the item we already have (but fetch body separately)
                        # Try to get body content
                        body_url = f"{CONFLUENCE_BASE_URL}/wiki/rest/api/content/{item_id}"
                        body_params = {"expand": "body.storage,body.view"}
                        body_response = requests.get(body_url, headers=headers, params=body_params)
                        if body_response.status_code == 200:
                            item_with_body = body_response.json()
                            all_pages.append(item_with_body)
                        else:
                            all_pages.append(item)
                        logging.info(f"Found page (partial): {item.get('title', 'N/A')}")
            
            # Check if there are more pages
            if len(results) < limit:
                break
            
            start += limit
            
        except requests.exceptions.RequestException as e:
            logging.error(f"Error fetching children: {e}")
            if hasattr(e, 'response') and e.response is not None:
                logging.error(f"Response status: {e.response.status_code}")
                logging.error(f"Response body: {e.response.text}")
            break
    
    return all_pages


def get_all_pages_in_folder(folder_id):
    """
    Gets all pages in a folder recursively.
    Uses recursive children fetching to get all pages and subfolders.
    Also tries CQL query as fallback.
    """
    # First, get the folder info
    folder_page = fetch_page_by_id(folder_id)
    if not folder_page:
        logging.error(f"Could not fetch folder {folder_id}")
        return [], "Unknown Folder", ""
    
    space_key = folder_page.get("space", {}).get("key", "")
    folder_title = folder_page.get("title", "Unknown Folder")
    
    logging.info(f"Fetching all pages in folder: {folder_title} (Space: {space_key})")
    
    # Try recursive approach first
    all_pages = get_children_recursive(folder_id)
    
    # If that didn't work, try CQL query approach
    if not all_pages:
        logging.info("Recursive approach found no pages, trying CQL query...")
        all_pages = get_pages_via_cql(folder_id, space_key, folder_title)
    
    return all_pages, folder_title, space_key


def get_pages_via_cql(folder_id, space_key, folder_title):
    """Try to get pages using CQL query with ancestor."""
    all_pages = []
    headers = get_confluence_auth_headers()
    url = f"{CONFLUENCE_BASE_URL}/wiki/rest/api/content/search"
    
    # Try different CQL queries
    cql_queries = [
        f'space = "{space_key}" AND type = page AND ancestor = {folder_id}',
        f'space = "{space_key}" AND type = page AND parent = {folder_id}',
        f'space = "{space_key}" AND type = page'
    ]
    
    for cql_query in cql_queries:
        logging.info(f"Trying CQL: {cql_query}")
        start = 0
        limit = 50
        
        while True:
            params = {
                "cql": cql_query,
                "limit": limit,
                "start": start,
                "expand": "body.storage,body.view,space,version,ancestors"
            }
            
            try:
                response = requests.get(url, headers=headers, params=params)
                response.raise_for_status()
                data = response.json()
                
                pages = data.get("results", [])
                if not pages:
                    break
                
                # Filter pages that are actually in this folder (check ancestors)
                for page in pages:
                    ancestors = page.get("ancestors", [])
                    ancestor_ids = [anc.get("id") for anc in ancestors if isinstance(anc, dict)]
                    if str(folder_id) in ancestor_ids or page.get("id") == str(folder_id):
                        all_pages.append(page)
                        logging.info(f"Found page via CQL: {page.get('title', 'N/A')}")
                
                if len(pages) < limit:
                    break
                
                start += limit
                
            except requests.exceptions.RequestException as e:
                logging.warning(f"CQL query failed: {e}")
                break
        
        if all_pages:
            break
    
    return all_pages


def parse_folder_url(url):
    """Extract folder ID from Confluence folder URL."""
    # Pattern: https://domain.atlassian.net/wiki/spaces/SPACEKEY/folder/FOLDERID
    pattern = r'/folder/(\d+)'
    match = re.search(pattern, url)
    if match:
        return match.group(1)
    
    # Alternative pattern: /pages/PAGEID or just the ID
    pattern2 = r'/pages/(\d+)'
    match2 = re.search(pattern2, url)
    if match2:
        return match2.group(1)
    
    raise ValueError(f"Could not extract folder ID from URL: {url}")


# --- EXPORT FUNCTIONS ---
def export_page_to_word(page, output_path):
    """Export a single page to Word document with tables."""
    doc = Document()
    
    # Add title
    title = page.get("title", "Untitled")
    heading = doc.add_heading(title, level=1)
    
    # Add metadata
    space = page.get("space", {})
    space_name = space.get("name", "") if isinstance(space, dict) else ""
    version = page.get("version", {})
    last_modified = version.get("when", "") if isinstance(version, dict) else ""
    version_by = version.get("by", {}) if isinstance(version, dict) else {}
    author = version_by.get("displayName", "") if isinstance(version_by, dict) else ""
    
    metadata_para = doc.add_paragraph()
    metadata_para.add_run(f"Space: {space_name}").bold = True
    metadata_para.add_run(f" | Author: {author} | Last Modified: {last_modified}")
    
    # Add URL
    page_url = f"{CONFLUENCE_BASE_URL}/wiki{page.get('_links', {}).get('webui', '')}"
    doc.add_paragraph(f"URL: {page_url}")
    doc.add_paragraph("")  # Empty line
    
    # Add content - try to extract with tables
    body_field = page.get("body", {})
    body_storage = body_field.get("storage", {})
    
    # If we have ADF format, extract with tables
    if isinstance(body_storage, dict) and ("type" in body_storage or "content" in body_storage):
        add_adf_content_to_word(doc, body_storage)
    else:
        # Fallback to plain text extraction
        body_text = parse_body_content(body_field, extract_plain_text=True)
        # Split into paragraphs and add
        paragraphs = body_text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())
    
    doc.save(output_path)
    return output_path


def add_adf_content_to_word(doc, node):
    """Recursively add ADF content to Word document, handling tables and expandable sections."""
    if not node or not isinstance(node, dict):
        return
    
    node_type = node.get("type", "")
    
    # Handle expand nodes (collapsible sections)
    if node_type == "expand":
        title = ""
        if "attrs" in node and "title" in node["attrs"]:
            title = node["attrs"]["title"]
        
        if title:
            doc.add_heading(title, level=2)
        
        if "content" in node and isinstance(node["content"], list):
            for child in node["content"]:
                add_adf_content_to_word(doc, child)
        return
    
    # Handle tables
    if node_type == "table":
        add_table_to_word(doc, node)
        return
    
    # Handle headings
    if node_type == "heading":
        level = node.get("attrs", {}).get("level", 1)
        text = extract_text_from_adf(node, include_tables=False)
        if text:
            doc.add_heading(text, level=level)
        return
    
    # Handle paragraphs
    if node_type == "paragraph":
        text = extract_text_from_adf(node, include_tables=False)
        if text.strip():
            doc.add_paragraph(text.strip())
        return
    
    # Handle bullet lists
    if node_type == "bulletList":
        if "content" in node and isinstance(node["content"], list):
            for list_item in node["content"]:
                if list_item.get("type") == "listItem":
                    item_text = extract_text_from_adf(list_item, include_tables=False)
                    if item_text.strip():
                        para = doc.add_paragraph(item_text.strip(), style='List Bullet')
        return
    
    # Handle numbered lists
    if node_type == "orderedList":
        if "content" in node and isinstance(node["content"], list):
            for idx, list_item in enumerate(node["content"], 1):
                if list_item.get("type") == "listItem":
                    item_text = extract_text_from_adf(list_item, include_tables=False)
                    if item_text.strip():
                        para = doc.add_paragraph(item_text.strip(), style='List Number')
        return
    
    # Recursively process content
    if "content" in node and isinstance(node["content"], list):
        for child in node["content"]:
            add_adf_content_to_word(doc, child)


def add_table_to_word(doc, table_node):
    """Add a table from ADF format to Word document."""
    if not table_node or table_node.get("type") != "table":
        return
    
    rows = []
    if "content" in table_node and isinstance(table_node["content"], list):
        for row_node in table_node["content"]:
            if row_node.get("type") == "tableRow":
                cells = []
                if "content" in row_node and isinstance(row_node["content"], list):
                    for cell_node in row_node["content"]:
                        if cell_node.get("type") in ["tableCell", "tableHeader"]:
                            cell_text = ""
                            if "content" in cell_node and isinstance(cell_node["content"], list):
                                for child in cell_node["content"]:
                                    child_text = extract_text_from_adf(child, include_tables=False)
                                    if child_text:
                                        cell_text += child_text + " "
                            cells.append(cell_text.strip())
                if cells:
                    rows.append(cells)
    
    if not rows:
        return
    
    # Create Word table
    max_cols = max(len(row) for row in rows) if rows else 0
    if max_cols == 0:
        return
    
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Light Grid Accent 1'
    
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx in range(max_cols):
            cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
            row.cells[col_idx].text = cell_text
            # Make header row bold
            if row_idx == 0:
                for paragraph in row.cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def export_page_to_text(page, output_path):
    """Export a single page to plain text file with tables."""
    with open(output_path, 'w', encoding='utf-8') as f:
        title = page.get("title", "Untitled")
        f.write("=" * 80 + "\n")
        f.write(f"{title}\n")
        f.write("=" * 80 + "\n\n")
        
        # Add metadata
        space = page.get("space", {})
        space_name = space.get("name", "") if isinstance(space, dict) else ""
        version = page.get("version", {})
        last_modified = version.get("when", "") if isinstance(version, dict) else ""
        version_by = version.get("by", {}) if isinstance(version, dict) else {}
        author = version_by.get("displayName", "") if isinstance(version_by, dict) else ""
        
        page_url = f"{CONFLUENCE_BASE_URL}/wiki{page.get('_links', {}).get('webui', '')}"
        
        f.write(f"Space: {space_name}\n")
        f.write(f"Author: {author}\n")
        f.write(f"Last Modified: {last_modified}\n")
        f.write(f"URL: {page_url}\n")
        f.write("\n")
        
        # Add content - use enhanced ADF extraction with tables
        body_field = page.get("body", {})
        body_storage = body_field.get("storage", {})
        
        # If we have ADF format, extract with tables
        if isinstance(body_storage, dict) and ("type" in body_storage or "content" in body_storage):
            body_text = extract_text_from_adf(body_storage, include_tables=True)
        else:
            # Fallback to plain text extraction
            body_text = parse_body_content(body_field, extract_plain_text=True)
        
        f.write(body_text)
        f.write("\n")
    
    return output_path


def process_folder(folder_url, output_dir):
    """Process a single folder and export all pages."""
    logging.info(f"\n{'='*60}")
    logging.info(f"Processing folder: {folder_url}")
    logging.info(f"{'='*60}")
    
    # Extract folder ID
    try:
        folder_id = parse_folder_url(folder_url)
        logging.info(f"Extracted folder ID: {folder_id}")
    except ValueError as e:
        logging.error(f"Error: {e}")
        return False
    
    # Get all pages in folder
    pages, folder_title, space_key = get_all_pages_in_folder(folder_id)
    
    if not pages:
        logging.warning(f"No pages found in folder: {folder_title}")
        return False
    
    logging.info(f"Found {len(pages)} pages in folder: {folder_title}")
    
    # Create output directory for this folder
    folder_output_dir = output_dir / sanitize_filename(folder_title)
    folder_output_dir.mkdir(parents=True, exist_ok=True)
    
    # Export each page
    success_count = 0
    for i, page in enumerate(pages, 1):
        page_title = page.get("title", f"Page_{i}")
        sanitized_title = sanitize_filename(page_title)
        
        logging.info(f"[{i}/{len(pages)}] Exporting: {page_title}")
        
        # Export to Word
        word_path = folder_output_dir / f"{sanitized_title}.docx"
        try:
            export_page_to_word(page, word_path)
            logging.info(f"  ✓ Word: {word_path.name}")
        except Exception as e:
            logging.error(f"  ✗ Word export failed: {e}")
        
        # Export to text
        txt_path = folder_output_dir / f"{sanitized_title}.txt"
        try:
            export_page_to_text(page, txt_path)
            logging.info(f"  ✓ Text: {txt_path.name}")
            success_count += 1
        except Exception as e:
            logging.error(f"  ✗ Text export failed: {e}")
    
    logging.info(f"\n✓ Completed folder '{folder_title}': {success_count}/{len(pages)} pages exported")
    logging.info(f"  Output directory: {folder_output_dir}")
    
    return True


def merge_folder_exports(folder_path, output_format="txt"):
    """
    Merge all exported files from a folder into one combined file.
    
    Args:
        folder_path: Path to the folder containing exported files
        output_format: 'txt' to merge text files, 'word' to merge Word documents
    """
    folder_path = Path(folder_path)
    if not folder_path.exists():
        logging.error(f"Folder not found: {folder_path}")
        return None
    
    if output_format == "txt":
        return merge_text_files(folder_path)
    elif output_format == "word":
        return merge_word_files(folder_path)
    else:
        logging.error(f"Unsupported format: {output_format}. Use 'txt' or 'word'")
        return None


def merge_text_files(folder_path):
    """Merge all .txt files from a folder into one combined text file."""
    txt_files = sorted(folder_path.glob("*.txt"))
    
    if not txt_files:
        logging.warning(f"No .txt files found in {folder_path}")
        return None
    
    output_file = folder_path / f"ALL_RELEASES_MERGED_{folder_path.name}.txt"
    
    logging.info(f"Merging {len(txt_files)} text files into {output_file.name}")
    
    with open(output_file, 'w', encoding='utf-8') as outfile:
        # Write header
        outfile.write("=" * 100 + "\n")
        outfile.write(f"ALL RELEASE NOTES - {folder_path.name.upper()}\n")
        outfile.write(f"Merged from {len(txt_files)} release notes\n")
        outfile.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        outfile.write("=" * 100 + "\n\n\n")
        
        for i, txt_file in enumerate(txt_files, 1):
            # Extract release name from filename
            release_name = txt_file.stem
            
            # Write separator
            outfile.write("\n" + "=" * 100 + "\n")
            outfile.write(f"RELEASE {i}/{len(txt_files)}: {release_name}\n")
            outfile.write("=" * 100 + "\n\n")
            
            # Read and write file content
            try:
                with open(txt_file, 'r', encoding='utf-8') as infile:
                    content = infile.read()
                    # Skip the header/metadata section (first 10 lines) to avoid duplication
                    lines = content.split('\n')
                    # Find where the actual content starts (after URL line)
                    content_start = 0
                    for idx, line in enumerate(lines):
                        if line.startswith('URL:'):
                            content_start = idx + 2  # Skip URL line and empty line
                            break
                    
                    # Write the content
                    outfile.write('\n'.join(lines[content_start:]))
                    outfile.write("\n\n")
                    
                logging.info(f"  ✓ Added: {release_name}")
            except Exception as e:
                logging.error(f"  ✗ Error reading {txt_file.name}: {e}")
                outfile.write(f"[ERROR: Could not read {txt_file.name}]\n\n")
        
        outfile.write("\n" + "=" * 100 + "\n")
        outfile.write(f"END OF MERGED RELEASE NOTES\n")
        outfile.write("=" * 100 + "\n")
    
    file_size = output_file.stat().st_size
    logging.info(f"✓ Merged file created: {output_file.name} ({file_size:,} bytes)")
    return output_file


def merge_word_files(folder_path):
    """Merge all .docx files from a folder into one combined Word document."""
    docx_files = sorted(folder_path.glob("*.docx"))
    
    if not docx_files:
        logging.warning(f"No .docx files found in {folder_path}")
        return None
    
    output_file = folder_path / f"ALL_RELEASES_MERGED_{folder_path.name}.docx"
    
    logging.info(f"Merging {len(docx_files)} Word documents into {output_file.name}")
    
    merged_doc = Document()
    
    # Add title page
    title = merged_doc.add_heading(f'ALL RELEASE NOTES - {folder_path.name.upper()}', 0)
    merged_doc.add_paragraph(f"Merged from {len(docx_files)} release notes")
    merged_doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    merged_doc.add_page_break()
    
    for i, docx_file in enumerate(docx_files, 1):
        release_name = docx_file.stem
        
        # Add release separator
        merged_doc.add_heading(f'RELEASE {i}/{len(docx_files)}: {release_name}', 1)
        merged_doc.add_paragraph("=" * 80)
        merged_doc.add_paragraph("")  # Empty line
        
        try:
            # Open source document
            source_doc = Document(docx_file)
            
            # Skip the first heading (title) and metadata, copy the rest
            content_started = False
            for para in source_doc.paragraphs:
                # Skip title and metadata paragraphs
                if not content_started:
                    if para.text.startswith("URL:"):
                        content_started = True
                        merged_doc.add_paragraph("")  # Empty line after URL
                        continue
                    if para.text.startswith("Space:") or para.text.startswith("Author:") or para.text.startswith("Last Modified:"):
                        continue
                    if para.style.name.startswith('Heading') and para.text:
                        # This is likely the title, skip it
                        continue
                    if not para.text.strip():
                        continue
                
                # Copy paragraph
                if content_started or para.text.strip():
                    new_para = merged_doc.add_paragraph()
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                    new_para.style = para.style
            
            # Copy tables
            for table in source_doc.tables:
                # Create new table with same dimensions
                new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                new_table.style = table.style
                
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        new_table.rows[row_idx].cells[col_idx].text = cell.text
                        # Copy formatting
                        for para in cell.paragraphs:
                            for run in para.runs:
                                new_run = new_table.rows[row_idx].cells[col_idx].paragraphs[0].add_run(run.text)
                                new_run.bold = run.bold
            
            logging.info(f"  ✓ Added: {release_name}")
        except Exception as e:
            logging.error(f"  ✗ Error reading {docx_file.name}: {e}")
            merged_doc.add_paragraph(f"[ERROR: Could not read {docx_file.name}]")
        
        # Add page break between releases (except for last one)
        if i < len(docx_files):
            merged_doc.add_page_break()
    
    # Add footer
    merged_doc.add_page_break()
    merged_doc.add_heading('END OF MERGED RELEASE NOTES', 1)
    
    merged_doc.save(output_file)
    file_size = output_file.stat().st_size
    logging.info(f"✓ Merged file created: {output_file.name} ({file_size:,} bytes)")
    return output_file


def main():
    """Main function."""
    # Check if this is a merge operation
    if len(sys.argv) >= 2 and sys.argv[1] == "--merge":
        if len(sys.argv) < 3:
            print("Usage: python confluence_folder_extractor.py --merge <folder_path> [txt|word]")
            print("\nExample:")
            print('  python confluence_folder_extractor.py --merge "confluence_exports/Deployment"')
            print('  python confluence_folder_extractor.py --merge "confluence_exports/Deployment" word')
            sys.exit(1)
        
        folder_path = Path(sys.argv[2])
        output_format = sys.argv[3] if len(sys.argv) > 3 else "txt"
        
        logging.info("=" * 60)
        logging.info("Confluence Folder Merger")
        logging.info("=" * 60)
        
        result = merge_folder_exports(folder_path, output_format)
        if result:
            logging.info(f"\n✓ Merge completed: {result}")
        else:
            logging.error("\n✗ Merge failed")
        return
    
    # Normal extraction mode
    if len(sys.argv) < 2:
        print("Usage: python confluence_folder_extractor.py <folder_url1> [folder_url2] ...")
        print("   OR: python confluence_folder_extractor.py --merge <folder_path> [txt|word]")
        print("\nExamples:")
        print('  python confluence_folder_extractor.py "https://your-domain.atlassian.net/wiki/spaces/YourSpace/folder/1234567890"')
        print('  python confluence_folder_extractor.py --merge "confluence_exports/Deployment"')
        print('  python confluence_folder_extractor.py --merge "confluence_exports/Deployment" word')
        sys.exit(1)
    
    folder_urls = sys.argv[1:]
    
    logging.info("=" * 60)
    logging.info("Confluence Folder Extractor")
    logging.info("=" * 60)
    logging.info(f"Processing {len(folder_urls)} folder(s)")
    logging.info(f"Output directory: {OUTPUT_BASE_DIR}")
    
    # Create output base directory
    OUTPUT_BASE_DIR.mkdir(parents=True, exist_ok=True)
    
    # Process each folder
    results = []
    for folder_url in folder_urls:
        success = process_folder(folder_url, OUTPUT_BASE_DIR)
        results.append((folder_url, success))
    
    # Summary
    logging.info("\n" + "=" * 60)
    logging.info("Summary")
    logging.info("=" * 60)
    for folder_url, success in results:
        status = "✓ Success" if success else "✗ Failed"
        logging.info(f"{status}: {folder_url}")
    
    successful = sum(1 for _, success in results if success)
    logging.info(f"\nTotal: {successful}/{len(results)} folders processed successfully")
    logging.info(f"Output location: {OUTPUT_BASE_DIR}")


if __name__ == "__main__":
    main()
